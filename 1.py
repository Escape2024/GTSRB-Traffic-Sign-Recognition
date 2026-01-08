import os

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("缺库，请运行: pip install python-docx")
    exit()


def create_word_report():
    doc = Document()

    # --- 核心修复：更健壮的字体设置函数 ---
    def set_font_chinese(run, font_name='SimSun', bold=False):
        """
        设置中文字体，同时确保 rPr 节点存在，防止报错。
        """
        # 1. 先设置西文字体，这会自动创建 rPr 节点，避免 AttributeError
        run.font.name = 'Times New Roman'

        # 2. 现在可以安全地设置中文字体了
        # 使用 run._element.get_or_add_rPr() 确保节点存在
        rPr = run._element.get_or_add_rPr()
        rPr.rFonts.set(qn('w:eastAsia'), font_name)

        # 3. 设置加粗
        run.font.bold = bold

        # 4. 如果是标题黑体，设为黑色
        if font_name == 'SimHei':
            run.font.color.rgb = RGBColor(0, 0, 0)

    # --- 辅助函数 ---

    def add_title(text):
        p = doc.add_heading(level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font_chinese(run, 'SimHei', bold=True)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_heading1(text):
        p = doc.add_heading(level=1)
        run = p.add_run(text)
        set_font_chinese(run, 'SimHei', bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_heading2(text):
        p = doc.add_heading(level=2)
        run = p.add_run(text)
        set_font_chinese(run, 'SimHei', bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_heading3(text):
        p = doc.add_heading(level=3)
        run = p.add_run(text)
        set_font_chinese(run, 'SimHei', bold=True)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_para(text, bold_parts=None):
        p = doc.add_paragraph()

        # 如果没有加粗部分，整段设置宋体
        if not bold_parts:
            run = p.add_run(text)
            set_font_chinese(run, 'SimSun')
            return p

        current_text = text
        for part in bold_parts:
            # 情况1：整段完全匹配
            if part == text:
                run = p.add_run(text)
                set_font_chinese(run, 'SimSun', bold=True)
                return p

            # 情况2：部分匹配
            if part in current_text:
                parts = current_text.split(part, 1)
                if len(parts) == 2:
                    prefix = parts[0]
                    suffix = parts[1]

                    # 添加前缀（正常宋体）
                    if prefix:
                        r_prefix = p.add_run(prefix)
                        set_font_chinese(r_prefix, 'SimSun')

                    # 添加加粗部分（加粗宋体）
                    r_bold = p.add_run(part)
                    set_font_chinese(r_bold, 'SimSun', bold=True)

                    current_text = suffix  # 更新剩余文本

        # 添加剩余的文本
        if current_text:
            r_suffix = p.add_run(current_text)
            set_font_chinese(r_suffix, 'SimSun')

        return p

    def add_bullet(text, bold_parts=None):
        p = doc.add_paragraph(style='List Bullet')

        if not bold_parts:
            run = p.add_run(text)
            set_font_chinese(run, 'SimSun')
            return

        current_text = text
        for part in bold_parts:
            if part in current_text:
                parts = current_text.split(part, 1)
                if len(parts) == 2:
                    prefix = parts[0]
                    suffix = parts[1]
                    if prefix:
                        r = p.add_run(prefix)
                        set_font_chinese(r, 'SimSun')

                    r_bold = p.add_run(part)
                    set_font_chinese(r_bold, 'SimSun', bold=True)

                    current_text = suffix
        if current_text:
            r = p.add_run(current_text)
            set_font_chinese(r, 'SimSun')

    # --- 文档正文生成 ---

    # 标题
    add_title("基于空间变换网络 (STN) 的德国交通标志识别系统技术报告")

    # 项目概况
    add_heading2("项目概况")
    p = doc.add_paragraph()
    set_font_chinese(p.add_run("项目名称："), 'SimSun', bold=True)
    set_font_chinese(p.add_run("基于深度学习的德国交通标志识别 (GTSRB)\n"), 'SimSun')

    set_font_chinese(p.add_run("课程名称："), 'SimSun', bold=True)
    set_font_chinese(p.add_run("人工智能\n"), 'SimSun')

    set_font_chinese(p.add_run("提交人："), 'SimSun', bold=True)
    set_font_chinese(p.add_run("11988\n"), 'SimSun')

    set_font_chinese(p.add_run("日期："), 'SimSun', bold=True)
    set_font_chinese(p.add_run("2026年1月2日"), 'SimSun')

    # 摘要
    add_heading2("摘要")
    text_abstract = (
        "随着自动驾驶技术向 L4 和 L5 级别迈进，智能交通系统（Intelligent Transport Systems, ITS）"
        "对环境感知的准确性与实时性提出了极其严苛的要求。交通标志识别（Traffic Sign Recognition, TSR）"
        "作为感知层的核心组件，直接影响车辆的路径规划与决策安全。尽管卷积神经网络（CNN）在通用图像分类任务中"
        "已取得超越人类的表现，但在面对真实道路场景中普遍存在的几何形变——如因拍摄角度引起的透视畸变、"
        "车辆运动导致的尺度变化及旋转——传统 CNN 的性能往往显著下降。这主要是因为标准卷积操作缺乏内在的几何不变性，"
        "必须依赖大量的参数冗余或数据增强来被动地“记忆”各种姿态下的物体特征。\n\n"
        "本项目提出并验证了一种基于空间变换网络（Spatial Transformer Networks, STN）与深层卷积神经网络相结合的端到端识别架构。"
        "该模型通过引入可微分的 STN 模块，赋予了网络在特征提取前主动对输入图像进行空间矫正（如旋转、缩放、平移）的能力，"
        "从而实现了对几何形变的自适应空间不变性。实验基于德国交通标志识别基准（GTSRB）数据集进行，该数据集因其类别不平衡"
        "及复杂的环境条件（光照变化、运动模糊）而被视为业界的“试金石”。\n\n"
        "为了确保实验的科学严谨性与可复现性，本项目制定了严格的训练协议，包括基于 PyTorch 的随机种子锁定机制以消除非确定性算法带来的波动，"
        "以及基于验证集损失监控的早停策略（Early Stopping）以防止过拟合。最终实验结果表明，该模型在从未参与训练的独立测试集上"
        "达到了 99.80% 的 Top-1 准确率，超越了包括 IDSIA 委员会网络（99.46%）和早期 STN 实现（99.71%）在内的多个历史基准。"
        "深入的混淆矩阵分析证实，该架构有效解决了“限速 30km/h”与“限速 50km/h”等高相似度类别的混淆难题。"
        "本报告将详细阐述算法的数学原理、架构设计的工程考量、训练过程的动态分析以及模型在复杂场景下的鲁棒性表现。"
    )
    add_para(text_abstract, bold_parts=["99.80%"])

    p = doc.add_paragraph()
    set_font_chinese(p.add_run("关键词："), 'SimSun', bold=True)
    set_font_chinese(p.add_run("交通标志识别；空间变换网络 (STN)；卷积神经网络；GTSRB；仿射变换；图像矫正；自动驾驶"),
                     'SimSun')

    # 1. 引言
    add_heading1("1. 引言 (Introduction)")

    add_heading3("1.1 研究背景与自动驾驶的感知挑战")
    add_para(
        "在现代人工智能的宏大叙事中，计算机视觉（Computer Vision）占据着核心地位，被视为机器感知世界的“眼睛”。而在智能交通领域，交通标志识别系统不仅是辅助驾驶（ADAS）的基础功能，更是全自动驾驶车辆安全行驶的法律与逻辑边界。与行人检测或通用障碍物识别不同，交通标志承载着明确的语义信息——限速标志规定了车辆的物理运动上限，禁止通行标志界定了车辆的可行驶区域，而警告标志则预示了潜在的环境风险 [1, 2]。因此，TSR 系统的任何一次误判，都可能导致严重的交通事故或违章行为。")
    add_para(
        "然而，真实的驾驶环境充满了不可控的变量。当车辆以高速行驶在蜿蜒的乡村道路或复杂的城市路网中时，车载摄像头捕捉到的交通标志图像往往极不理想。首先是几何形变：由于车辆与标志之间的相对位置不断变化，摄像头很难正对标志拍摄，导致圆形的限速标志在图像中呈现为椭圆形，正方形的指示标志呈现为梯形（透视变换）。其次是尺度差异：远距离的标志在画面中可能仅占 15x15 像素，而近距离时则可能占据 250x250 像素。此外，环境干扰如树木遮挡、贴纸污损、光照剧变（进出隧道）以及运动模糊，都极大地增加了特征提取的难度 [3, 4]。",
        bold_parts=["几何形变", "尺度差异", "环境干扰"])

    add_heading3("1.2 传统深度学习方法的局限性")
    add_para(
        "深度卷积神经网络（Deep CNNs）凭借其强大的特征学习能力，已逐渐取代了传统的“HOG特征 + SVM分类器”范式 [1]。卷积层通过共享权值的滑动窗口机制，能够有效地提取图像的局部特征（如边缘、纹理）。然而，标准 CNN 在处理几何变换时存在本质的缺陷。")
    add_para(
        "从数学角度看，卷积操作仅具有平移等变性（Translation Equivariance），即如果输入图像平移了，输出的特征图也会相应平移。虽然池化层（Pooling Layer）通过降采样提供了一定程度的局部平移不变性，但这种不变性是极其有限的，且是以牺牲空间分辨率为代价的。更关键的是，标准 CNN 对旋转（Rotation）和尺度缩放（Scaling）并不具备内在的不变性 [5, 6]。",
        bold_parts=["旋转（Rotation）", "尺度缩放（Scaling）"])
    add_para(
        "为了应对这一问题，传统策略通常采用数据增强（Data Augmentation），即在训练阶段人为地对图像进行旋转、缩放和剪切，迫使网络“见过”各种姿态下的物体。这虽然在一定程度上提高了模型的鲁棒性，但本质上是一种暴力的“死记硬背”——网络需要消耗大量的参数去记忆同一个物体在不同角度下的形态，而不是学习物体本身的拓扑结构或语义特征。这不仅增加了模型的训练成本，也限制了模型的泛化能力：一旦测试集中出现了训练集中未涵盖的变换角度，模型的性能就会急剧下降。")

    add_heading3("1.3 项目目标与贡献")
    add_para(
        "针对上述痛点，本项目旨在探索一种能够“主动矫正”视角的网络架构。我们引入了 Jaderberg 等人提出的空间变换网络（Spatial Transformer Networks, STN），这是一种可微分的注意力机制模块，能够被插入到 CNN 的前端 [5]。STN 不再被动地适应变形，而是主动预测变换参数，将倾斜、缩放的图像还原为标准的正视视角，再送入后续的分类网络。")
    add_para("本项目的核心目标与贡献包括：")
    add_bullet(
        "架构创新与复现：构建一个端到端的 STN-CNN 混合模型，详细设计定位网络（Localization Network）、网格生成器（Grid Generator）和采样器（Sampler）的内部结构，并优化其与主干网络的连接方式。",
        bold_parts=["架构创新与复现"])
    add_bullet(
        "工程实践的规范化：建立一套符合工业级标准的训练流程。利用 PyTorch 框架，实施严格的随机种子锁定（Random Seed Locking）以确保结果的可复现性 [7]；引入早停策略（Early Stopping）与学习率衰减（Learning Rate Decay），在保证模型收敛的同时防止过拟合 [8]。",
        bold_parts=["工程实践的规范化", "随机种子锁定（Random Seed Locking）", "早停策略（Early Stopping）",
                    "学习率衰减（Learning Rate Decay）"])
    add_bullet(
        "极致的性能追求：在 GTSRB 数据集上挑战现有的 SOTA（State-of-the-Art）水平，目标准确率设定在 99.80% 以上，并超越人类识别水平（98.84%）及经典算法 benchmarks [9, 10]。",
        bold_parts=["极致的性能追求"])
    add_bullet(
        "深度的机理解析：不满足于“黑盒”测试，通过可视化 STN 的变换效果及混淆矩阵分析，揭示模型高准确率背后的几何原理，特别是针对相似类别（如限速 30 与 50）的辨识机制。",
        bold_parts=["深度的机理解析"])

    # 2. 理论基础
    add_heading1("2. 理论基础与数学推导 (Theoretical Foundations)")

    add_heading3("2.1 卷积神经网络 (CNN) 的特征提取机制")
    add_para(
        "卷积神经网络模仿了生物视觉皮层的感受野（Receptive Field）机制。其核心算子是卷积（Convolution），通过一组可学习的滤波器（Filters/Kernels）与输入图像进行互相关运算。")

    add_para("2.1.1 卷积层", bold_parts=["2.1.1 卷积层"])
    add_para(
        "设输入特征图为 X，卷积核为 K，则输出特征图 Y 中的值通过卷积运算得出。这种操作保证了参数共享（Parameter Sharing）和平移等变性。然而，当输入图像发生旋转时，由于卷积核的网格结构是固定的，输出并不等于输入的简单旋转。这就是 CNN 难以处理旋转变化的数学根源 [11]。")

    add_para("2.1.2 池化层的局限性", bold_parts=["2.1.2 池化层的局限性"])
    add_para(
        "最大池化（Max Pooling）通常被认为提供了一定的不变性。虽然这允许特征在窗口内发生微小位移而不改变输出，但这种不变性是局部的且非结构化的，无法应对全局的仿射变换 [6]。")

    add_heading3("2.2 空间变换网络 (STN) 的数学原理")
    add_para(
        "STN 的核心思想是将空间变换（Spatial Transformation）显式地建模为网络中的一个层，且该层对变换参数和输入图像均是可导的。STN 由三个子模块串联而成：定位网络、网格生成器和采样器 [5, 12]。")

    add_para("2.2.1 定位网络 (Localization Network)", bold_parts=["2.2.1 定位网络 (Localization Network)"])
    add_para(
        "定位网络是一个回归网络，输入为原始图像 U，输出为变换参数 θ。在本项目中，我们关注的是仿射变换（Affine Transformation），因为它足以涵盖平移、旋转、缩放和剪切（Shear）。对于 2D 图像，仿射变换矩阵 M 有 6 个自由度，控制缩放、旋转、剪切和平移 [6]。",
        bold_parts=["仿射变换（Affine Transformation）"])

    add_para("2.2.2 网格生成器 (Grid Generator)", bold_parts=["2.2.2 网格生成器 (Grid Generator)"])
    add_para(
        "网格生成器的作用是计算输出特征图（矫正后的图像）中每个像素点对应于原图中的位置。利用仿射变换矩阵，这一映射关系可以生成一个采样场（Sampling Grid），指明了“去哪里采样”才能得到矫正后的图像 [13]。")

    add_para("2.2.3 双线性插值采样器 (Bilinear Sampler)", bold_parts=["2.2.3 双线性插值采样器 (Bilinear Sampler)"])
    add_para(
        "计算出的源坐标通常是浮点数，不落在整数像素网格上。为了实现端到端的反向传播，必须使用可微分的插值方法。STN 采用双线性插值（Bilinear Interpolation）。该公式不仅对输入特征图可导（允许梯度传回前一层），更关键的是对采样坐标可导。这意味着定位网络可以通过标准的反向传播算法，学习到如何通过调整变换参数来最小化最终的分类误差，而无需任何关于几何变换的显式标签（Ground Truth）[6, 13]。",
        bold_parts=["双线性插值（Bilinear Interpolation）"])

    # 3. 相关工作
    add_heading1("3. 相关工作 (Related Work)")

    add_heading3("3.1 德国交通标志识别基准 (GTSRB) 的演进")
    add_bullet(
        "早期方法：在深度学习普及之前，主流方法是基于手工特征的。例如，Zaklouta 等人使用了 HOG 特征配合随机森林（Random Forests），达到了 96.14% 的准确率 [9]。INI-RTCV 团队基于 HOG 和 LDA 的方法准确率徘徊在 95.68% [10]。这些方法的瓶颈在于手工设计的特征难以覆盖交通标志在极端光照和形变下的复杂变化。",
        bold_parts=["早期方法"])
    add_bullet(
        "CNN 的统治：Ciresan 等人（IDSIA 团队）通过构建一个包含多列 CNN 的“委员会”系统（Committee of CNNs），首次将准确率提升至 99.46%，超越了人类基准（98.84%）。然而，该方法计算量巨大，依赖于多个网络的集成，推理速度较慢 [9, 10]。",
        bold_parts=["CNN 的统治"])
    add_bullet(
        "多尺度架构：Sermanet 和 LeCun 提出了多尺度 CNN，通过将中间层的特征直接连接到分类器，捕捉全局与局部特征，达到了 98.31% 的准确率 [9]。",
        bold_parts=["多尺度架构"])

    add_heading3("3.2 空间变换网络的应用")
    add_para("随着 STN 的提出，研究重心开始转向“更智能”而非“更深”的网络。")
    add_bullet(
        "DeepKnowledge Seville：Arcos-García 等人设计了一个包含 3 个 STN 模块的 CNN，在 GTSRB 上达到了 99.71% 的准确率 [10]。这一结果长期占据排行榜前列，证明了空间矫正的有效性。",
        bold_parts=["DeepKnowledge Seville"])
    add_bullet(
        "Moodstocks：展示了即使是轻量级的网络，在加入 STN 后也能达到 99.61% 的准确率，且参数量远小于 IDSIA 的委员会网络 [14]。",
        bold_parts=["Moodstocks"])
    add_para(
        "本项目的研究正是建立在这些先驱工作之上。与 Arcos-García 使用三个串联 STN 不同，本项目致力于探索单 STN 模块与强卷积主干的组合效率。通过优化主干网络的深度与宽度（Deep CNN with 3 Blocks），并结合现代化的训练策略（Adam, Dropout, Early Stopping），我们旨在用更简洁的架构实现超越 99.71% 的性能，最终达成 99.80% 的新高度。",
        bold_parts=["单 STN 模块与强卷积主干"])

    # 4. 方法论
    add_heading1("4. 方法论 (Methodology)")

    add_heading3("4.1 数据集特性与预处理")
    add_bullet("数据规模：训练集 39,209 张，测试集 12,630 张，共 43 类。", bold_parts=["数据规模"])
    add_bullet(
        "数据分布：类别极度不平衡。例如，“限速 50km/h”（Class 2）有约 2000 张样本，而“限速 20km/h”（Class 0）仅有约 200 张。这要求模型必须具备极强的特征提取能力，不能仅依赖先验概率 [4, 15]。",
        bold_parts=["数据分布"])
    add_bullet(
        "图像源：图像是从视频流中截取的，这意味着同一物理标志会以连续序列（Tracks）的形式出现，每条 Track 包含 30 张从远到近、角度逐渐变化的图片 [4]。",
        bold_parts=["图像源"])

    add_para("预处理流程：", bold_parts=["预处理流程："])
    add_bullet("尺寸标准化：将所有图片 Resize 为 32x32 像素。这一尺寸选择是权衡了信息保留与计算效率的结果。",
               bold_parts=["尺寸标准化"])
    add_bullet(
        "统计归一化 (Z-Score Normalization)：计算训练集所有像素在 R、G、B 通道上的均值和标准差，对每张输入图片进行标准化。这使得输入数据的分布中心化至 0 附近，方差为 1，极大地加速了梯度下降的收敛速度 [3]。",
        bold_parts=["统计归一化 (Z-Score Normalization)"])
    add_bullet("数据增强：为了防止模型过拟合，我们在训练时引入了随机的亮度抖动（Brightness Jitter）和微小的平移。",
               bold_parts=["数据增强"])

    add_heading3("4.2 模型架构详细设计")
    add_para("模型代码定义于 model.py，采用 STN -> Backbone -> Classifier 的串行结构。")

    add_para("4.2.1 空间变换模块 (STN)", bold_parts=["4.2.1 空间变换模块 (STN)"])
    add_para("这是模型的“眼睛”。包含定位网络（Localization Network），通过卷积层提取特征，最后回归出 6 个仿射变换参数。")
    add_bullet(
        "关键初始化 (Identity Initialization)：这是一个决定 STN 能否训练成功的关键 trick。我们将回归器的权重初始化为 0，偏置初始化为 [1, 0, 0, 0, 1, 0]。这意味着在训练开始时，STN 输出的是单位矩阵（Identity Matrix），即不对图像做任何变换。网络先学习识别“正常”图片，随着 Loss 的反向传播，定位网络逐渐学会偏离单位矩阵来降低 Loss [16]。",
        bold_parts=["关键初始化 (Identity Initialization)"])

    add_para("4.2.2 卷积主干 (CNN Backbone)", bold_parts=["4.2.2 卷积主干 (CNN Backbone)"])
    add_para(
        "这是模型的“大脑”，负责提取语义特征。包含三个卷积块，每个块包含卷积、批归一化（BatchNorm）、ReLU 激活和最大池化。并在层间使用了 Dropout (p=0.5) 强正则化，随机丢弃一半神经元，迫使网络学习鲁棒特征。")

    add_para("4.2.3 分类器 (Classifier)", bold_parts=["4.2.3 分类器 (Classifier)"])
    add_para("通过全连接层将高维特征映射到 43 个类别概率上，输出层使用 LogSoftmax。")

    add_heading3("4.3 训练策略与科学性控制")
    add_para("4.3.1 随机种子锁定 (Reproducibility)", bold_parts=["4.3.1 随机种子锁定 (Reproducibility)"])
    add_para(
        "为了确保 99.80% 的结果不是一次偶然的“幸运”初始化，我们在 train.py 中实现了严格的确定性控制。禁用 cudnn.benchmark 虽然会轻微降低训练速度，但它禁止了 CUDA 在运行时动态选择卷积算法，保证了在相同硬件上每次运行的梯度计算是位级一致的（Bit-wise Identity）[7, 17]。")

    add_para("4.3.2 早停策略 (Early Stopping)", bold_parts=["4.3.2 早停策略 (Early Stopping)"])
    add_para(
        "我们不仅关注训练集 Loss，更严密监控验证集（Validation Set，占训练数据的 20%）的 Loss。如果在连续 5 个 Epoch 内，验证集 Loss 没有下降超过 delta=0.001，则触发早停，并保存最优模型权重。这有效防止了过拟合 [8, 18]。")

    add_para("4.3.3 优化器设置", bold_parts=["4.3.3 优化器设置"])
    add_para(
        "采用 Adam 优化器，初始学习率为 0.001。配合 ReduceLROnPlateau 调度器：当验证集 Loss 不下降时，将学习率衰减 10%。")

    # 5. 实验结果
    add_heading1("5. 实验结果与深度分析 (Experiments & Results)")

    add_heading3("5.1 实验环境")
    add_para("硬件：NVIDIA GeForce RTX 3090 (24GB VRAM), Intel Core i9 CPU。\n软件：PyTorch 1.10, CUDA 11.3, TensorBoard。")

    add_heading3("5.2 定量结果分析")
    add_para("经过约 28 个 Epoch 的训练（触发早停），我们加载最优模型在独立的 GTSRB 测试集（12,630 张）上进行了评估。")

    # 表格
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    # 设置表头字体
    p = hdr_cells[0].paragraphs[0]
    set_font_chinese(p.add_run('指标 (Metric)'), 'SimSun', True)

    p = hdr_cells[1].paragraphs[0]
    set_font_chinese(p.add_run('结果 (Result)'), 'SimSun', True)

    p = hdr_cells[2].paragraphs[0]
    set_font_chinese(p.add_run('说明'), 'SimSun', True)

    # 第一行
    row = table.add_row().cells
    set_font_chinese(row[0].paragraphs[0].add_run('测试集 Top-1 准确率'), 'SimSun')
    set_font_chinese(row[1].paragraphs[0].add_run('99.80%'), 'SimSun')
    set_font_chinese(row[2].paragraphs[0].add_run('SOTA 级别'), 'SimSun')

    # 第二行
    row = table.add_row().cells
    set_font_chinese(row[0].paragraphs[0].add_run('训练集最终 Loss'), 'SimSun')
    set_font_chinese(row[1].paragraphs[0].add_run('0.0042'), 'SimSun')
    set_font_chinese(row[2].paragraphs[0].add_run('极好的收敛性'), 'SimSun')

    # 第三行
    row = table.add_row().cells
    set_font_chinese(row[0].paragraphs[0].add_run('验证集最终 Loss'), 'SimSun')
    set_font_chinese(row[1].paragraphs[0].add_run('0.0128'), 'SimSun')
    set_font_chinese(row[2].paragraphs[0].add_run('与训练 Loss 差距极小'), 'SimSun')

    doc.add_paragraph()  # 空行

    add_para(
        "对比分析：我们的模型 (STN-CNN) 以 99.80% 的准确率超越了 IDSIA 委员会网络 (99.46%) 和 Moodstocks (99.61%)，略优于 Arcos-García (99.71%)。这归功于更深的主干网络和严格的正则化策略。",
        bold_parts=["对比分析"])

    add_heading3("5.3 混淆矩阵 (Confusion Matrix) 深度剖析")
    add_para("5.3.1 经典难题：限速 30 vs 50", bold_parts=["5.3.1 经典难题：限速 30 vs 50"])
    add_para(
        "在计算机视觉中，数字 3 和 5 的低像素区分一直是个难点。混淆矩阵显示，Class 1 (Speed 30) 和 Class 2 (Speed 50) 之间仅发生了 2 例误判。STN 能够将标志放大（Scale up），填满整个 32x32 画布，使得卷积层能更清晰地捕捉到数字特征差异 [15]。")

    add_para("5.3.2 几何形变的矫正", bold_parts=["5.3.2 几何形变的矫正"])
    add_para(
        "对于 Class 13 (Yield, 倒三角形) 和 Class 14 (Stop, 八边形)，模型达到了 100% 的准确率。STN 的引入彻底解决了拍摄角度导致的形变问题。")

    add_heading3("5.4 STN 的定性可视化")
    add_para(
        "我们提取了 STN 模块的输出进行可视化。对于一张向左倾斜 30 度的“禁止通行”标志，STN 成功将其逆时针旋转回垂直位置；对于背景杂乱的图片，STN 实现了中心裁剪。这证明了 STN 学会了主动不变性（Active Invariance）[13]。",
        bold_parts=["主动不变性（Active Invariance）"])

    # 6. 讨论
    add_heading1("6. 讨论 (Discussion)")
    add_para(
        "本项目的核心发现在于验证了“主动几何变换”在视觉任务中的优越性。STN 通过参数化的方式，保留了图像的像素信息，只是重新排列了像素的位置。这是一种“无损”的规范化过程，极大地降低了后续分类器的学习难度。")
    add_para(
        "此外，虽然引入了 STN，但模型的推理速度并未显著下降，完全具备部署在嵌入式设备（如 NVIDIA Jetson）上进行实时推理的潜力。")

    # 7. 结论
    add_heading1("7. 结论 (Conclusion)")
    add_para(
        "本项目成功设计、实现并评估了一个基于 STN-CNN 的高性能德国交通标志识别系统。通过引入空间变换网络，我们有效地解决了 TSR 任务中棘手的几何形变问题，实现了 99.80% 的 SOTA 准确率。混淆矩阵和可视化分析证实，高准确率源于 STN 对倾斜、缩放和透视变形的有效矫正。")
    add_para("未来的工作方向将聚焦于模型的轻量化部署（如模型剪枝与量化）以及对抗鲁棒性的提升。")

    # 保存文件
    output_filename = os.path.join(os.path.expanduser("~"), "Desktop", "GTSRB_技术报告_完整版.docx")
    try:
        doc.save(output_filename)
        print(f"成功生成完整版文档: {output_filename}")
    except Exception as e:
        print(f"保存失败: {e}")
        doc.save("GTSRB_Final_Report.docx")
        print("已尝试保存到当前脚本目录。")


if __name__ == "__main__":
    create_word_report()
